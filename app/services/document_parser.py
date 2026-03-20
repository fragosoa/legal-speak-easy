from dataclasses import dataclass
from io import BytesIO

import pdfplumber
from docx import Document

from app.core.exceptions import ExtractionError, UnsupportedFileTypeError

# Magic bytes for file type detection
_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"

_MIN_TEXT_LENGTH = 100


@dataclass
class ParsedDocument:
    text: str
    page_count: int
    word_count: int
    file_type: str
    extraction_method: str


def parse(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Detect file type by magic bytes and extract text accordingly."""
    if file_bytes[:4] == _PDF_MAGIC:
        return _parse_pdf(file_bytes, filename)
    elif file_bytes[:4] == _DOCX_MAGIC:
        return _parse_docx(file_bytes, filename)
    else:
        raise UnsupportedFileTypeError(
            "Unsupported file type. Only PDF and DOCX files are accepted."
        )


def _parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    pages_text: list[str] = []

    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                pages_text.append(text)

    full_text = "\n\n".join(pages_text).strip()

    if len(full_text) < _MIN_TEXT_LENGTH:
        raise ExtractionError(
            "Could not extract readable text from this document. "
            "The file may be a scanned image. Please provide a text-based PDF."
        )

    return ParsedDocument(
        text=full_text,
        page_count=page_count,
        word_count=len(full_text.split()),
        file_type="pdf",
        extraction_method="pdfplumber",
    )


def _parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    doc = Document(BytesIO(file_bytes))
    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Extract table cells (common in work contracts for salary/benefits)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in parts:
                    parts.append(cell_text)

    full_text = "\n\n".join(parts).strip()

    if len(full_text) < _MIN_TEXT_LENGTH:
        raise ExtractionError(
            "Could not extract readable text from this DOCX file. "
            "The document appears to be empty or contains only images."
        )

    # DOCX doesn't have a native page count concept; estimate from word count
    word_count = len(full_text.split())
    estimated_pages = max(1, word_count // 300)

    return ParsedDocument(
        text=full_text,
        page_count=estimated_pages,
        word_count=word_count,
        file_type="docx",
        extraction_method="python-docx",
    )
